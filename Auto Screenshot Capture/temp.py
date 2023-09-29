# -*- coding: utf-8 -*-
"""
Spyder Editor

This is a temporary script file.
"""

#Created by Poojith Reddy

"""Plan to Design Actions
1.Taking screenshots  from backend.
2.Compare and removing the Screenshots.
3.Insert into Word File.
"""

#----Importing Modules
import numpy as np
import cv2
import pyautogui
import time as t
import matplotlib.pyplot as plt
import docx
from docx.shared import Inches


#comparing two adjacent images and removing if any dupliacates
def image_compare(image, image2):
    original = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    contrast = cv2.cvtColor(image2, cv2.COLOR_BGR2GRAY)
    
    # the 'Mean Squared Error' between the two images is the
    	# sum of the squared difference between the two images;
    	# NOTE: the two images must have the same dimension
    err = np.sum((original.astype("float") - contrast.astype("float")) ** 2)
    err /= float(original.shape[0] * original.shape[1])
    	
    	# return the MSE, the lower the error, the more "similar"
    	# the two images are
    return(err)



doc = docx.Document()
# Add a Title to the document
doc.add_heading('Testing Document', 0)
t1=t.time()
image = pyautogui.screenshot()
image = cv2.cvtColor(np.array(image),cv2.COLOR_RGB2BGR)
for i in range(10):
    t.sleep(1)
    # take screenshot using pyautogui
    image2 = pyautogui.screenshot()
       
    # since the pyautogui takes as a 
    # PIL(pillow) and in RGB we need to 
    # convert it to numpy array and BGR 
    # so we can write it to the disk
    image2 = cv2.cvtColor(np.array(image2),
                         cv2.COLOR_RGB2BGR)   
    if(image_compare(image, image2)>10):
        # writing it to the disk using opencv
        cv2.imwrite("screenshot/image"+str(i)+"1.png", image2)
        doc.add_picture("screenshot/image"+str(i)+"1.png",width=Inches(7), height=Inches(3))
    image = image2
doc.save('Captured_screenshots_file.docx')

    