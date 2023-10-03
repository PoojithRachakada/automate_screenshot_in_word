# -*- coding: utf-8 -*-
"""
Created on Sun Oct  1 09:16:10 2023

@author: poojith
"""

from tkinter import *
import tkinter as tk
from tkinter import filedialog
import numpy as np
import cv2
import pyautogui
import docx
from docx.shared import Inches
import os

filename = ""
doc = docx.Document()
full_details={}
current_details={'Screenshot_comments':[],'Total_Screenshots':0}
current_total_screenshots=0


#showing history values in new window
class Table:     
    def __init__(self,root):         
        # code for creating table
        # for i in range(total_rows):
        #     for j in range(total_columns):
                 
        #         self.e = Entry(root, width=20, fg='blue',
        #                        font=('Arial',16,'bold'))
                 
        #         self.e.grid(row=i, column=j)
        #         self.e.insert(END, lst[i][j])
        total_rows = len(full_details)
        list_full_details = list(full_details.items())
        #Heading Details
        self.e =tk.Label(root,text="File Name",width=20, borderwidth=3, relief="raised",bg='darkgreen',font=('Arial',10,'bold'))
        self.e.grid(row=0,column=0)
        self.e =tk.Label(root,text="Doc Title",width=30, borderwidth=3, relief="raised", bg='darkgreen',font=('Arial',10,'bold'))
        self.e.grid(row=0,column=1)
        self.e =tk.Label(root,text="T.Screenshots",width=10, borderwidth=3, relief="raised", bg='darkgreen',font=('Arial',10,'bold'))
        self.e.grid(row=0,column=2)
        self.e =tk.Label(root,text="File Path",  borderwidth=3, relief="raised",bg='darkgreen',font=('Arial',10,'bold'))
        self.e.grid(row=0,column=3)
        # self.e = Label(root,text="Screenshot comments", borderwidth=3, relief="raised", bg='darkgreen',font=('Arial',10,'bold'))
        # self.e.grid(row=0,column=4)
        
        for i in range(total_rows):
            for j in range(4):     
                # self.e = Entry(root, width=20, fg='blue',font=('Arial',8))
                self.e = tk.Label(root,text="File Path",borderwidth=3, relief="sunken",font=('Arial',8,'bold'))
                if(i%2==0):
                    self.e.config(fg='deepskyblue')
                else:
                    self.e.config(fg='indigo')
                self.e.grid(row=i+1,column=j)
                if(j==0):
                    # temp=list_full_details[i][j]
                    # self.e.insert(END,list_full_details[i][j])
                    self.e.config(text=list_full_details[i][j],width=20)
                elif(j==1):    
                    temp=list_full_details[i][1]
                    # print("temp-",temp['heading'])
                    # self.e.insert(END,temp['heading'])
                    self.e.config(text=temp['heading'],width=30)
                elif(j==2):
                    temp=list_full_details[i][1]
                    # self.e.insert(END,temp['Total_Screenshots'])
                    self.e.config(text=temp['Total_Screenshots'],width=10)
                elif(j==3):
                    temp=list_full_details[i][1]
                    # self.e.insert(END,temp['path'])
                    self.e.config(text=temp['path'])
                # elif(j==4):
                #     temp=list_full_details[i][1]
                #     # self.e.insert(END,temp['Screenshot_comments'])
                #     self.e.config(text=temp['Screenshot_comments'])
                


def add_head_doc(t_heading):
    main_heading = entry.get()
    global doc
    doc.add_heading(main_heading, t_heading)

def add_picture_doc(image):
    add_head_doc(3)
    #storing current details
    current_details['Total_Screenshots']+=1
    current_details['Screenshot_comments'].append(entry.get())
    image = np.ascontiguousarray(image)
    # writing it to the disk using opencv
    cv2.imwrite("screenshot/image1.png", image)
    doc.add_picture("screenshot/image1.png",width=Inches(7), height=Inches(3))
    os.remove("screenshot/image1.png")
    

def take_screenshot():
    win.wm_attributes("-alpha",0.0)#Hiding overlay screen while taking screenshot
    # take screenshot using pyautogui
    image = pyautogui.screenshot()
    win.wm_attributes("-alpha",1.0)
    add_picture_doc(image)
    
    
def sel_folder():
    global filename
    # File save dialog.
    filename =filedialog.asksaveasfilename()#filedialog.askdirectory()
    # print(filename)

def reset_to_new():
    but2.place_forget()
    but3.place_forget()
    but5.place_forget()
    but1.place(x=180,y=60)
    label.config(text="Doc Title")
    current_details.clear()     #clearing the current store details
    current_details['Total_Screenshots']=0
    current_details['Screenshot_comments']=[]
    
def save_file():
    sel_folder()
    print(filename)    
    global doc
    doc.save(filename+".docx")
    doc = docx.Document()    
    #storing current data to Full data
    global current_total_screenshots
    # global current_screenshot_comments
    current_details['path']=filename
    full_details[filename.split('/')[-1]]=current_details.copy()
    print(full_details)
    reset_to_new()

    
def save_doc_heading():   
    add_head_doc(0)    
    current_details['heading']=entry.get() #storing data
    entry.delete(0, END)
    but1.place_forget()
    label.config(text="Screenshot text")
    # but2.pack(side='left')
    but2.place(x=180,y=60) #Screenshot button
    # but3.pack(side='right') 
    but3.place(x=360,y=60) #save button
    but5.place(x=50,y=60) #refresh button

def history_details():
    #pass #https://www.geeksforgeeks.org/create-table-using-tkinter/
    h_root = tk.Tk()
    t = Table(h_root)
    h_root.mainloop()

# Create an instance of tkinter frame or window
win = tk.Tk()

# Set the size of the window
win.geometry("500x100")
win.wm_attributes("-alpha",1.0) #transperancy

canvas = tk.Canvas(win, bg="darkturquoise", width=600, height=30)
canvas.create_text(150, 10, text="Overlay Screenshot capture", font=('', 13))
canvas.pack()

#Initialize a Label to display the User Input
label=tk.Label(win, text="Doc Title", font=("Courier 12 bold"),bg="darkslategray1",height=1)
label.place(x=0,y=35)
# label.pack(padx=10,pady=20)

#Create an Entry widget to accept User Input
entry= tk.Entry(win, width= 55, bg='bisque4' ,fg='white')
entry.place(x=155,y=35)
# entry.focus_set()
# entry.pack(padx=20,pady=20)
# Add a Button to start/stop the loop

but1 = tk.Button(win, text="Save Title", bg="lightpink1" ,command=lambda:save_doc_heading()) #Button for saving Heading,,,,,command=lambda:save_doc_heading() 
but1.place(x=180,y=60)

but2 = tk.Button(win, text="Screenshot", bg="goldenrod2" ,command=take_screenshot) #Button for taking Screenshot 
# but2.pack(side='left')
but2.place_forget()

but3 = tk.Button(win, text="Save File" ,command=save_file) #Button for Saving the file
# but3.pack(side='right')
but3.place_forget()

but4 = tk.Button(win, text="History",bg='limegreen',fg='white' ,command=history_details) #History Button
but4.place(x=400,y=5)

but5 = tk.Button(win, text="Reset",bg='purple1',fg='white' ,command=reset_to_new) #Reset Button
but5.place_forget()

win.attributes('-topmost',True) #keeping the screen on top of all windows
win.config(bg='darkslategray1')
win.resizable(False,False)
win.mainloop()